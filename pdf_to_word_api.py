import unittest
from flask import Flask, request, send_file, Response
from unittest.mock import MagicMock
from docx import Document, text
from docx.shared import RGBColor
import os
import csv
from datetime import datetime
import io, json
import pdfplumber
import requests
import tempfile
from pdf2image import convert_from_path  # Added for PDF to image conversion
import base64

app = Flask(__name__)

# Azure OpenAI configuration
endpoint = "https://openai-eus-ti-poc-shared-resources.openai.azure.com/"
api_key = "f6f65288e4ec4f89a92387cb877c4b17"  # Replace with your actual API key
deployment_id = "chat-gpt-4o"  # or "chat-gpt-4o-mini"
api_version = "2025-01-01-preview"

# We'll use requests directly instead of the OpenAI client


def extract_text_from_pdf(pdf_path):
    """Extracts text directly from a PDF using pdfplumber."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            all_text = ""
            for page in pdf.pages:
                all_text += page.extract_text() + "\n"
            return all_text
    except Exception as e:
        print(f"PDF text extraction error: {e}")
        return ""


def call_azure_openai(messages, deployment_id, api_key, endpoint, api_version):
    """Call Azure OpenAI API directly using requests."""
    url = f"{endpoint}openai/deployments/{deployment_id}/chat/completions?api-version={api_version}"
    headers = {
        "Content-Type": "application/json",
        "api-key": api_key
    }
    
    payload = {
        "messages": messages,
        "temperature": 0,
        "top_p": 0.95,
        "frequency_penalty": 0,
        "presence_penalty": 0
    }
    
    response = requests.post(url, headers=headers, json=payload)
    
    if response.status_code == 200:
        return response.json()
    else:
        print(f"Error calling Azure OpenAI API: {response.status_code}, {response.text}")
        raise Exception(f"Azure OpenAI API error: {response.status_code}")


def extract_changes_from_pdf(pdf_path, _, deployment_id, pdf_filename):
    """
    Extracts tracked changes from a PDF by converting it to images and using Azure OpenAI vision capabilities.
    This works better for PDFs that contain tracked changes which may not be properly extracted as text.
    """
    try:
        print(f"Processing PDF: {pdf_filename}")
        # Convert PDF to images
        images = convert_pdf_to_images(pdf_path)
        
        if not images:
            print(f"Failed to convert PDF to images: {pdf_filename}")
            return [], 0
        
        print(f"Successfully converted PDF to {len(images)} images")
        
        # Process images with Azure OpenAI
        return process_images_with_azure_openai(images, deployment_id, pdf_filename)
        
    except Exception as e:
        print(f"Error extracting changes from PDF: {e}")
        return [], 0


def convert_pdf_to_images(pdf_path):
    """Converts PDF pages to images."""
    try:
        # For Windows, you may need to specify the path to poppler
        # Attempt to use the default path first
        try:
            images = convert_from_path(
                pdf_path,
                dpi=200,  # Adjust DPI for quality vs. performance
                fmt="png"
            )
        except Exception as poppler_error:
            print(f"Standard conversion failed, attempting with poppler path: {poppler_error}")
            # Try with explicit poppler path - adjust this path as needed
            poppler_path = r"C:\Users\JD15806\Code\poppler-24.08.0\Library\bin"
            images = convert_from_path(
                pdf_path,
                dpi=200,
                fmt="png",
                poppler_path=poppler_path
            )
        
        print(f"Converted PDF to {len(images)} images")
        return images
    except Exception as e:
        print(f"Error converting PDF to images: {e}")
        return []


def process_images_with_azure_openai(images, deployment_id, pdf_filename):
    """Processes images with Azure OpenAI vision capabilities to extract tracked changes."""
    extracted_changes = []
    total_token_usage = 0
    
    # Process images in batches to avoid exceeding token limits
    batch_size = 4  # Adjust based on token usage and performance
    
    for i in range(0, len(images), batch_size):
        batch_images = images[i:i+batch_size]
        batch_changes, batch_tokens = process_image_batch(batch_images, i, deployment_id)
        
        extracted_changes.extend(batch_changes)
        total_token_usage += batch_tokens
        
        print(f"Processed batch {i//batch_size + 1}/{(len(images) + batch_size - 1)//batch_size}, pages {i+1}-{min(i+batch_size, len(images))}")
    
    return extracted_changes, total_token_usage


def process_image_batch(images, start_page, deployment_id):
    """Process a batch of images with Azure OpenAI."""
    # Convert images to base64
    base64_images = []
    for img in images:
        buffered = io.BytesIO()
        img.save(buffered, format="PNG")
        img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
        base64_images.append(img_base64)
    
    # Prepare prompt
    system_message = """You are an expert document editor analyzing PDF pages generated from Word documents with track changes.
    Your task is to identify and extract only the paragraphs that have been modified with track changes. 
    Focus exclusively on changes such as insertions, deletions, and replacements.
    Do not extract paragraphs without modification."""
    
    user_content = """Please identify paragraphs with tracked changes in these PDF pages.
    - Only consider paragraphs that begin with a numerical prefix (e.g., "1.", "2.1", "3.a").
    - For paragraphs, represent formatting changes as follows:
        - Underlined text: <u>text</u>
        - Strikethrough text: <s>text</s>
        - Highlighted text: <highlight>text</highlight>
    - Do not return the paragraphs that have no change.
    
    Return the output in JSON format with each element containing:
    - 'paragraph_number': The paragraph number
    - 'content': The paragraph content with tracked changes marked using the specified formatting tags
    """
    
    # Create message content with images
    messages = [
        {"role": "system", "content": system_message},
        {
            "role": "user", 
            "content": [
                {"type": "text", "text": user_content}
            ] + [{"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img}"}} for img in base64_images]
        }
    ]
    
    # Call Azure OpenAI API
    url = f"{endpoint}openai/deployments/{deployment_id}/chat/completions?api-version={api_version}"
    headers = {
        "Content-Type": "application/json",
        "api-key": api_key
    }
    
    payload = {
        "messages": messages,
        "temperature": 0,
        "top_p": 0.95,
        "max_tokens": 4000,
        "stream": False
    }
    
    try:
        response = requests.post(url, headers=headers, json=payload)
        
        if response.status_code == 200:
            response_json = response.json()
            response_content = response_json['choices'][0]['message']['content'].strip()
            
            # Create output directory if it doesn't exist
            output_dir = os.path.join(os.path.dirname(__file__), 'outputs')
            os.makedirs(output_dir, exist_ok=True)
            
            # Save raw response for debugging
            debug_file = os.path.join(output_dir, f'raw_response_batch_{start_page+1}.txt')
            with open(debug_file, 'w', encoding='utf-8') as f:
                f.write(response_content)
            
            try:
                # Clean up response content
                if response_content.startswith('```json'):
                    # Extract content between triple backticks
                    content_start = response_content.find('[')
                    content_end = response_content.rfind(']') + 1
                    if content_start >= 0 and content_end > content_start:
                        response_content = response_content[content_start:content_end]
                
                # Parse the JSON response
                batch_changes = json.loads(response_content)
                
                # Add page numbers to changes
                for change in batch_changes:
                    if isinstance(change, dict) and 'paragraph_number' in change:
                        change['page'] = f"Pages {start_page+1}-{start_page+len(images)}"
                
                # Save processed changes
                output_file = os.path.join(output_dir, f'changes_batch_{start_page+1}.json')
                with open(output_file, 'w', encoding='utf-8') as f:
                    json.dump(batch_changes, f, indent=2)
                
                token_usage = response_json.get('usage', {}).get('total_tokens', 0)
                print(f"Successfully processed batch {start_page+1} with {len(batch_changes)} changes")
                return batch_changes, token_usage
                
            except json.JSONDecodeError as e:
                print(f"Failed to parse JSON response: {e}")
                print(f"Response content after cleaning: {response_content}")
                return [], 0
        else:
            print(f"Error calling Azure OpenAI API: {response.status_code}, {response.text}")
            return [], 0
    except Exception as e:
        print(f"Exception while calling Azure OpenAI API: {e}")
        return [], 0


def add_formatted_text(paragraph, text):
    """Adds text to a paragraph, applying formatting markers."""
    if text == "":
        return
    
    start_underline = text.find("<u>")
    start_strikethrough = text.find("<s>")
    start_highlight = text.find("<highlight>")

    if start_underline != -1:
        run = paragraph.add_run(text[0:start_underline])
        run_underline = paragraph.add_run(text[start_underline+3:text.find("</u>")])
        run_underline.underline = True
        add_formatted_text(paragraph,text[text.find("</u>")+4:])
    elif start_strikethrough != -1:
        run = paragraph.add_run(text[0:start_strikethrough])
        run_strikethrough = paragraph.add_run(text[start_strikethrough+3:text.find("</s>")])
        run_strikethrough.font.strike = True
        add_formatted_text(paragraph,text[text.find("</s>")+4:])
    elif start_highlight != -1:
        run = paragraph.add_run(text[0:start_highlight])
        run_highlight = paragraph.add_run(text[start_highlight+10:text.find("</highlight>")])
        run_highlight.font.highlight_color = 4  # wdYellow
        add_formatted_text(paragraph,text[text.find("</highlight>")+11:])
    else:
        paragraph.add_run(text)


def fill_word_template(template_path, changes):
    """Fills a Word template with extracted changes.
       Replaces {{txtNo}} with paragraph_number and {{txtParagraph}} with content.
    """
    doc = Document(template_path)

     # Find the target table
    target_table = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{txtNo}}" in cell.text:
                    target_table = table
                    break
            if target_table:
                break
        if target_table:
            break

    if not target_table:
        raise ValueError("Template table with {{txtNo}} placeholder not found")

    # Get template row (the one with placeholders)
    template_row = None
    template_row_index = -1
    for i, row in enumerate(target_table.rows):
        for cell in row.cells:
            if "{{txtNo}}" in cell.text:
                template_row = row
                template_row_index = i
                break
        if template_row:
            break

    # Delete all rows below template row
    for _ in range(len(target_table.rows) - template_row_index - 1):
        target_table._element.remove(target_table.rows[-1]._element)

    # Add rows for each change
    for change in changes:
        new_row = target_table.add_row()
        
        # Copy cell formatting from template row
        for i, template_cell in enumerate(template_row.cells):
            new_cell = new_row.cells[i]
            # Copy cell properties if needed
            new_cell._tc.get_or_add_tcPr().append(template_cell._tc.get_or_add_tcPr())

        # Fill data
        for i, cell in enumerate(new_row.cells):
            if i == 0:  # First column - paragraph number
                cell.text = str(change.get('paragraph_number', ''))
            elif i == 1:  # Second column - content
                # Clear existing paragraph to avoid duplicate text
                cell.paragraphs[0].clear()
                # Add formatted text
                add_formatted_text(cell.paragraphs[0], change.get('content', ''))

    # Remove template row
    target_table._element.remove(template_row._element)

    return doc


def log_api_call(pdf_filename, word_filename, api_info, token_usage=None):
    """Logs API call information to a CSV file."""
    log_data = {
        "timestamp": datetime.now().isoformat(),
        "pdf_filename": pdf_filename,
        "word_filename": word_filename,
        "api_info": api_info,
        "token_usage": token_usage,
    }
    with open("api_log.csv", "a", newline="", encoding="utf-8") as csvfile:
        fieldnames = ["timestamp", "pdf_filename", "word_filename", "api_info", "token_usage"]
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        if csvfile.tell() == 0:
            writer.writeheader()
        writer.writerow(log_data)


@app.route("/")
@app.route("/convert", methods=["POST"])
def convert_pdf_to_word():
    """API endpoint to convert a PDF with tracked changes to a Word document."""
    if "file" not in request.files or "template" not in request.files:
        return Response("Please provide both a PDF file and a Word template.", status=400)

    pdf_file = request.files["file"]
    template_file = request.files["template"]

    if not pdf_file.filename.endswith(".pdf") or not template_file.filename.endswith(".docx"):
        return Response("Invalid file types. Please provide a PDF file and a Word template.", status=400)

    pdf_filename = pdf_file.filename
    template_filename = template_file.filename

    # Create temporary directory for this request
    import tempfile
    temp_dir = tempfile.mkdtemp()
    
    pdf_path = os.path.join(temp_dir, pdf_filename)
    template_path = os.path.join(temp_dir, template_filename)
    pdf_file.save(pdf_path)
    template_file.save(template_path)
    
    output_path = None
    response = None
    
    try:
        print(f"Processing PDF with image-based extraction: {pdf_filename}")
        # Use the new image-based extraction method
        changes, total_token_usage = extract_changes_from_pdf(pdf_path, None, deployment_id, pdf_filename)
        
        # Save total changes for testing
        output_dir = os.path.join(os.path.dirname(__file__), 'outputs')
        os.makedirs(output_dir, exist_ok=True)
        total_changes_file = os.path.join(output_dir, f'total_changes_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json')
        with open(total_changes_file, 'w', encoding='utf-8') as f:
            json.dump({
                'pdf_filename': pdf_filename,
                'total_changes': len(changes),
                'changes': changes,
                'token_usage': total_token_usage
            }, f, indent=2)

        if not changes:
            print("No changes detected with image-based extraction, trying fallback text extraction")
            # If image-based extraction fails or finds no changes, try fallback with direct text extraction
            text_content = extract_text_from_pdf(pdf_path)
            
            # Prepare messages for Azure OpenAI
            system_message = """You are an expert document editor. You are given text extracted from a PDF file that was generated from a Word document with track changes enabled.
            
            Your task is to identify and extract only the paragraphs that have been modified or deleted by track changes. Please focus exclusively on changes such as insertions, deletions, and replacements. Do not extract paragraphs without modification."""
            
            user_content = f"""Please identify paragraphs with tracked changes in this PDF text.
            - Only consider paragraphs that begin with a numerical prefix (e.g., "1.", "2.1", "3.a").
            - For paragraphs, represent formatting changes as follows:
                - Underlined text: <u>text</u>
                - Strikethrough text: <s>text</s>
                - Highlighted text: <highlight>text</highlight>
            - Do not return the paragraphs that have no change.
            
            Return the output in JSON format with each element containing:
            - 'paragraph_number': The paragraph number
            - 'content': The paragraph content with tracked changes marked using the specified formatting tags
            
            Here is the text:
            {text_content}
            """
            
            messages = [
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_content}
            ]
            
            # Call Azure OpenAI with text
            url = f"{endpoint}openai/deployments/{deployment_id}/chat/completions?api-version={api_version}"
            headers = {
                "Content-Type": "application/json",
                "api-key": api_key
            }
            
            payload = {
                "messages": messages,
                "temperature": 0,
                "top_p": 0.95,
                "max_tokens": 4000,
                "stream": False
            }
            
            text_response = requests.post(url, headers=headers, json=payload)
            
            if text_response.status_code == 200:
                text_response_json = text_response.json()
                text_content_response = text_response_json['choices'][0]['message']['content'].strip()
                
                try:
                    # Try to parse the JSON response
                    text_changes = json.loads(text_content_response)
                    
                    # Use these changes if we found some
                    if text_changes:
                        changes = text_changes
                        total_token_usage = text_response_json.get('usage', {}).get('total_tokens', 0)
                        print(f"Found {len(changes)} changes using text-based extraction")
                except json.JSONDecodeError as e:
                    print(f"Failed to parse JSON response from text-based extraction: {e}")

        if not changes:
            return Response("No changes detected in the PDF document.", status=400)

        print(f"Processing {len(changes)} extracted changes")
        output_doc = fill_word_template(template_path, changes)

        # Create result directory if it doesn't exist
        result_dir = os.path.join(os.path.dirname(__file__), 'result')
        os.makedirs(result_dir, exist_ok=True)

        # Generate unique filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        word_filename = f"output_{os.path.splitext(pdf_filename)[0]}_{timestamp}.docx"
        result_path = os.path.join(result_dir, word_filename)
        
        # Save file to result directory
        output_doc.save(result_path)

        log_api_call(pdf_filename, word_filename, f"Azure OpenAI API: {deployment_id}", total_token_usage)

        # Return JSON response with filename and token usage
        response_data = {
            "filename": word_filename,
            "token_usage": total_token_usage,
            "file_path": result_path
        }
        
        return Response(json.dumps(response_data), mimetype='application/json')

    except Exception as e:
        print(f"Error: {e}")
        return Response(f"An error occurred: {e}", status=500)

    finally:
        # Clean up files safely
        try:
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            if os.path.exists(template_path):
                os.remove(template_path)
            if output_path and os.path.exists(output_path):
                os.remove(output_path)
            # Remove temp directory
            if os.path.exists(temp_dir):
                os.rmdir(temp_dir)
        except Exception as e:
            print(f"Cleanup error: {e}")
            # Continue even if cleanup fails


if __name__ == "__main__":
    if not os.path.exists("api_log.csv"):
        with open("api_log.csv", "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["timestamp", "pdf_filename", "word_filename", "api_info", "token_usage"])
    app.run(debug=True, host="0.0.0.0", port=5000)